import streamlit as st
from crewai import Agent, Task, Crew, Process
from crewai import LLM
import tempfile
from typing import List, Dict
import os
import docx
import PyPDF2
from io import BytesIO
import time
from io import BytesIO


def create_processing_indicator():
    return st.empty()

# Function to read PDF files
def read_pdf(file):
    try:
        if isinstance(file, str):
            # If it's a string (file path), open the file first
            with open(file, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
        else:
            # If it's already a file object, reset the pointer to the beginning
            file.seek(0)
            pdf_reader = PyPDF2.PdfReader(file)
        
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        raise ValueError(f"Error reading PDF: {str(e)}")

def read_docx(file):
    try:
        if isinstance(file, str):
            doc = docx.Document(file)
        else:
            # Reset file pointer to beginning
            file.seek(0)
            doc = docx.Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        raise ValueError(f"Error reading DOCX: {str(e)}")

def read_vtt(file):
    try:
        if isinstance(file, str):
            with open(file, 'r', encoding='utf-8') as f:
                content = f.read()
        else:
            # Reset file pointer to beginning
            file.seek(0)
            content = file.read().decode('utf-8')
        
        # Remove VTT header and metadata
        lines = content.split('\n')
        text = ""
        skip_line = True
        
        for line in lines:
            # Skip WEBVTT header and timestamps
            if line.strip() == 'WEBVTT':
                continue
            if '-->' in line:
                continue
            if not line.strip():
                continue
            if line.strip().isdigit():
                continue
            # Add the actual text content
            text += line.strip() + " "
        
        return text.strip()
    except Exception as e:
        raise ValueError(f"Error reading VTT: {str(e)}")

def read_txt(file):
    try:
        if isinstance(file, str):
            with open(file, 'rb') as f:
                return f.read().decode('utf-8')
        else:
            # Reset file pointer to beginning
            file.seek(0)
            return file.read().decode('utf-8')
    except Exception as e:
        raise ValueError(f"Error reading TXT: {str(e)}")

def read_file(file):
    try:
        if isinstance(file, str):
            file_type = file.split('.')[-1].lower()
        else:
            file_type = file.name.split('.')[-1].lower()
        
        handlers = {
            'pdf': read_pdf,
            'docx': read_docx,
            'txt': read_txt,
            'vtt': read_vtt
        }
        
        if file_type not in handlers:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        return handlers[file_type](file)
    except Exception as e:
        raise ValueError(f"Error processing file: {str(e)}")



class TranscriptProcessor:
    def __init__(self, api_key):
        self.llm = LLM(model="gemini/gemini-2.0-flash", api_key=api_key)
        self.setup_agents()

    def setup_agents(self):
        self.content_analyzer = Agent(
            role="Content and Structure Analyzer",
            goal="Analyze transcript content and create structured document with proper formatting",
            backstory="Expert at analyzing academic content and creating well-structured documents",
            llm=self.llm,
            verbose=True
        )

        self.quote_extractor = Agent(
            role="Quote and Insight Extractor",
            goal="Extract and categorize meaningful quotes and key insights",
            backstory="Specialist in identifying impactful quotes and critical insights from academic discussions. Extract the exact quotes, not paraphrased.",
            llm=self.llm,
            verbose=True
        )

        self.content_writer = Agent(
            role="Content Writer and Organizer",
            goal="""Create a comprehensive document with the following sections in order:
                1. Title and Speaker Information
                2. Key Quotes
                3. Closing Statements
                4. Briefing Document
                5. Key Themes and Ideas
                6. Notable Quotes with Context
                7. FAQ Section (exactly 10)
                8. Quiz Questions (exactly 10)
                9. Quiz Answer Key
                10. Essay Questions (5-7)
                11. Essay Answers
                12. Speaker Bio""",
            backstory="Experienced in creating detailed academic content and educational materials with strict adherence to formatting and section ordering",
            llm=self.llm,
            verbose=True
        )

    def process_transcript(self, transcript_text: str, speaker_name: str, progress_bar, status_text) -> str:
        try:
            # Initialize progress
            progress_bar.progress(0)
            status_text.text("Initializing analysis...")
            time.sleep(1)

            # Task 1: Initial Analysis
            progress_bar.progress(0.1)
            status_text.text("Analyzing content structure...")
            time.sleep(1)


            # Task 1: Initial Analysis and Title/Quote Formation
            analysis_task = Task(
                description=f"""Analyze this transcript: {transcript_text}

                Based on this exact transcript content, analyze and create:
                1. Title and speaker section:
                - Extract or formulate an appropriate title for the talk
                - Include speaker name: {speaker_name}
                - Include any mentioned affiliation (excluding any pharmaceutical company references like Pfizer)

                2. Key Quotes (exactly 20-25):
                - Extract the MOST IMPACTFUL and INSIGHTFUL quotes from the transcript
                - Use the exact quotes verbatim, improve the language to make it formal and professional
                - Choose only quotes that represent key technical insights, profound thoughts, or critical information
                - Prioritize quotes that demonstrate speaker expertise and deep domain knowledge
                - Ensure each quote has clear relevance to the main topics discussed
                - Format as numbered list with quotation marks
                - Exclude filler content, generic statements, or mundane remarks
                - IMPORTANT: Exclude any quotes containing references to pharmaceutical companies
                - Remove or redact any pharmaceutical company mentions
                - Preserving the original quote make the language formal and professional

                3. Closing Statements:
                - Identify and extract the speaker's concluding remarks
                - Include any final thoughts or takeaways mentioned
                
                IMPORTANT: DO NOT INCLUDE ANY PHARMACEUTICAL COMPANY REFERENCES OR MENTIONS IN THE OUTPUT.
                IMPORTANT: Ensure no other names are mentioned in the file apart from the speaker name provided or anything related to Pfizer.""",
                expected_output="""Structured markdown sections with title, speaker info, KEY and IMPACTFUL numbered quotes, and closing statements.
                Ensure no other names are mentioned in the file apart from the speaker name provided or anything related to Pfizer.""",
                agent=self.content_analyzer
            )
            
            # Update progress for first task completion
            progress_bar.progress(0.25)
            status_text.text("Extracting key quotes and themes...")
            time.sleep(1)

            # Task 2: Extract Themes and Create Briefing
            quotes_task = Task(
                description=f"""Using this transcript: {transcript_text}

                And the previous analysis, create:
                1. Create a comprehensive briefing document:
                - Summarize main discussion points
                - Identify key arguments and insights
                - Structure the information logically

                2. Extract and organize key themes and ideas:
                - List major topics discussed
                - Provide supporting evidence from transcript
                - Connect related concepts
                - Remove any pharmaceutical company mentions or references

                3. Highlight significant quotes:
                - Select most impactful statements
                - Provide context for each quote
                - Explain significance
                - Exclude quotes containing pharmaceutical company references

                4. Write a detailed conclusion:
                - Summarize key takeaways
                - Connect main themes
                - Highlight implications
                IMPORTANT: DO NOT INCLUDE ANY PHARMACEUTICAL COMPANY REFERENCES OR MENTIONS IN THE OUTPUT.
                IMPORTANT: Ensure no other names are mentioned in the file apart from the speaker name provided or anything related to Pfizer.""",
                expected_output="Detailed markdown sections with briefing, themes, quotes, and conclusion. Ensure no other names are mentioned in the file apart from the speaker name provided or anything related to Pfizer.",
                agent=self.quote_extractor
            )
            
            # Update progress for second task completion
            progress_bar.progress(0.45)
            status_text.text("Creating educational content...")
            time.sleep(1)


            # Task 3: Create Educational Content
            content_task = Task(
                description=f"""Using this transcript: {transcript_text}

                Create a comprehensive document with the following sections IN THIS EXACT ORDER:

                # Title and Speaker Information
                - Formulate an appropriate title without name being Pfizer centric or any pharmaceutical company references
                - Include speaker name: {speaker_name}
                - Include affiliation (excluding any pharmaceutical company references)

                # Key Quotes
                - identify around 20-25 key quotes
                - Extract ONLY the most significant and insightful exact quotes from the transcript
                - Choose quotes that demonstrate the speaker's expertise and deep knowledge
                - Select quotes that reveal important technical insights or profound thoughts
                - Number each quote
                - Include brief context where relevant
                - Avoid generic statements or filler content
                - Preserving the original quote make the language formal and professional
                - Let the quotes be 2-3 sentences long.

                # Closing Statements
                - Extract the speaker's concluding remarks
                - Include final thoughts and takeaways
                - Around 100 words

                # Briefing Document
                - Summarize main discussion points
                - Identify key arguments and insights
                - Structure information logically

                # Key Themes and Ideas
                - List major topics discussed
                - Provide supporting evidence
                - Connect related concepts

                # Notable Quotes with Context
                - Present significant statements
                - Explain their importance
                - Connect to main themes

                # FAQ Section (around 20-25)
                - Create relevant questions
                - Provide 2-3 sentence answers
                - Cover main topics

                # Quiz Questions (around 15-20)
                - Mix of multiple choice and short answer
                - Base on transcript content
                - Include key concepts

                # Quiz Answer Key
                - Provide detailed explanations
                - Reference transcript
                - Explain reasoning

                # Essay Questions (5-7)
                - Create thought-provoking questions
                - Focus on main themes
                - Include response guidance

                # Essay Answers
                - For each essay question, provide a detailed and well-structured essay answer.
                - Each essay answer should have a title that is the essay question itself.
                - Structure each essay answer into three paragraphs.
                - Reference the transcript to support your points.
                - Explain your reasoning clearly and logically.
                - Maintain a proper format and professional tone throughout the essay answers.

                # Speaker Bio
                - Create bio for {speaker_name}
                - Include mentioned information
                - Focus on expertise

                IMPORTANT: Maintain this exact section order and use proper markdown formatting with clear section headers.
                Also make sure no other names are mentioned in the file apart from the speaker name provided anything related to Pfizer.
                2. DO NOT INCLUDE ANY PHARMACEUTICAL COMPANY REFERENCES IN ANY SECTION
                3. Remove or redact any industry-specific company mentions
                4. Focus on academic and technical content only
                5. Exclude any industry affiliations or relationships
                6. Keep content generic where industry references would normally appear.
                7. Where ever the quote are presented make sure they are presented in a formal and professional manner.""",
                expected_output="""A comprehensive markdown document with all sections in the specified order:
                    Title and Speaker Information, Key Quotes, Closing Statements, Briefing Document,
                    Key Themes and Ideas, Notable Quotes, FAQ Section, Quiz Questions, Quiz Answer Key,
                    Essay Questions, Essay Answers, and Speaker Bio. and also make sure no other names are mentioned in the file 
                    apart from the speaker name provided anything related to Pfizer.""",
                agent=self.content_writer
            )
            

            # Create Crew
            crew = Crew(
                agents=[self.content_analyzer, self.quote_extractor, self.content_writer],
                tasks=[analysis_task, quotes_task, content_task],
                process=Process.sequential,
                verbose=True
            )

            # Execute crew with intermittent progress updates
            status_text.text("Processing with AI agents...")
            
            # Simulate progress during processing
            for i in range(46, 95, 5):
                progress_bar.progress(i/100)
                time.sleep(2)  # Adjust this value based on your actual processing time
                st.spinner(f"Processing... {i}% complete")
                status_text.text(f"Processing... {i}% complete")

            result = crew.kickoff()

            # Post-process to remove any remaining pharmaceutical references
            processed_result = str(result).replace("Pfizer", "[REDACTED]")
            processed_result = processed_result.replace("pfizer", "[REDACTED]")
        

            # Final progress update
            progress_bar.progress(1.0)
            status_text.text("Processing complete!")
            time.sleep(1)
            
            return str(result)

        except Exception as e:
            raise Exception(f"Error in processing: {str(e)}")

def markdown_to_pdf(markdown_content: str, filename: str):
    from xhtml2pdf import pisa
    import markdown2
    
    # Convert markdown to HTML with extra features
    html_content = markdown2.markdown(
        markdown_content,
        extras=[
            "tables",
            "break-on-newline",
            "cuddled-lists",
            "fenced-code-blocks"
        ]
    )
    
    # Add CSS styling
    styled_html = f"""
    <html>
    <head>
        <style>
            @page {{
                size: a4 portrait;
                margin: 2cm;
            }}
            body {{
                font-family: Arial, sans-serif;
                font-size: 12px;
                line-height: 1.6;
            }}
            h1 {{
                font-size: 24px;
                color: #2c3e50;
                margin-top: 20px;
            }}
            h2 {{
                font-size: 20px;
                color: #2c3e50;
                margin-top: 15px;
            }}
            h3 {{
                font-size: 16px;
                color: #2c3e50;
            }}
            blockquote {{
                margin: 10px 0;
                padding-left: 10px;
                border-left: 3px solid #2c3e50;
                color: #666;
            }}
            code {{
                background-color: #f5f5f5;
                padding: 2px 5px;
                border-radius: 3px;
            }}
            ul, ol {{
                margin: 10px 0;
                padding-left: 20px;
            }}
            p {{
                margin: 10px 0;
            }}
        </style>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """
    
    # Create PDF in memory
    pdf_data = BytesIO()
    pisa.CreatePDF(styled_html, dest=pdf_data)
    pdf_data.seek(0)
    return pdf_data

def markdown_to_docx(markdown_content: str) -> BytesIO:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import markdown2
    
    # Create a new Document
    doc = Document()
    
    # Define styles
    styles = {
        'title': {'size': Pt(24), 'bold': True, 'color': RGBColor(44, 62, 80), 'space_after': Pt(12)},
        'section_header': {'size': Pt(18), 'bold': True, 'color': RGBColor(44, 62, 80), 'space_before': Pt(18), 'space_after': Pt(8)},
        'subsection_header': {'size': Pt(14), 'bold': True, 'color': RGBColor(52, 73, 94), 'space_before': Pt(14), 'space_after': Pt(6)},
        'normal_text': {'size': Pt(11), 'space_after': Pt(6)},
        'quote_text': {'size': Pt(11), 'italic': True, 'left_indent': Inches(0.5), 'space_after': Pt(8)}
    }
    
    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Split content by headers and paragraphs
    sections = markdown_content.split('\n#')
    
    # Process first section (if no # at start)
    if not markdown_content.startswith('#'):
        if sections[0].strip():
            p = doc.add_paragraph(sections[0].strip())
            p.style = doc.styles['Normal']
            p.paragraph_format.space_after = styles['normal_text']['space_after']
        sections = sections[1:]
    
    # Process each section
    for section in sections:
        if section.strip():
            # Split header from content
            parts = section.split('\n', 1)
            header = parts[0].strip('# ')
            
            # Add header as a main section
            h = doc.add_heading(header, level=1)
            font = h.runs[0].font
            font.size = styles['section_header']['size']
            font.bold = styles['section_header']['bold']
            font.color.rgb = styles['section_header']['color']
            h.paragraph_format.space_before = styles['section_header']['space_before']
            h.paragraph_format.space_after = styles['section_header']['space_after']
            
            # Add content if there is any
            if len(parts) > 1:
                content = parts[1].strip()
                lines = content.split('\n')
                
                i = 0
                while i < len(lines):
                    line = lines[i].strip()
                    
                    # Skip empty lines
                    if not line:
                        i += 1
                        continue
                    
                    # Check for subsection headers (##)
                    if line.startswith('##'):
                        subheader = line.strip('#').strip()
                        h2 = doc.add_heading(subheader, level=2)
                        font = h2.runs[0].font
                        font.size = styles['subsection_header']['size']
                        font.bold = styles['subsection_header']['bold']
                        font.color.rgb = styles['subsection_header']['color']
                        h2.paragraph_format.space_before = styles['subsection_header']['space_before']
                        h2.paragraph_format.space_after = styles['subsection_header']['space_after']
                    
                    # Check for lists
                    elif line.startswith('- ') or line.startswith('* '):
                        p = doc.add_paragraph(line[2:], style='List Bullet')
                        p.paragraph_format.space_after = styles['normal_text']['space_after']
                        
                    elif line.startswith('1.') or line.startswith('1)') or (line[0].isdigit() and line[1:].startswith('. ')):
                        # Extract the number and the content
                        num_end = line.find('.')
                        if num_end == -1:
                            num_end = line.find(')')
                        
                        if num_end != -1:
                            content_start = num_end + 1
                            quote_content = line[content_start:].strip()
                            
                            # Check if this is a quote in the Key Quotes section
                            if header == "Key Quotes" or "Numbered Quotes" in header:
                                p = doc.add_paragraph()
                                p.paragraph_format.left_indent = styles['quote_text']['left_indent']
                                p.paragraph_format.space_after = styles['quote_text']['space_after']
                                
                                # Add the number
                                number_run = p.add_run(f"{line[:content_start]} ")
                                number_run.bold = True
                                
                                # Add the quote
                                quote_run = p.add_run(quote_content)
                                quote_run.italic = True
                                
                            else:
                                p = doc.add_paragraph(line, style='List Number')
                                p.paragraph_format.space_after = styles['normal_text']['space_after']
                    
                    # Regular paragraphs
                    else:
                        p = doc.add_paragraph(line)
                        p.paragraph_format.space_after = styles['normal_text']['space_after']
                        
                        # If this is a potential title at the beginning of the document
                        if header == "Title and Speaker Information" and i == 0:
                            for run in p.runs:
                                run.font.size = styles['title']['size']
                                run.font.bold = styles['title']['bold']
                                run.font.color.rgb = styles['title']['color']
                            p.paragraph_format.space_after = styles['title']['space_after']
                            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    i += 1
    
    # Save to memory
    docx_data = BytesIO()
    doc.save(docx_data)
    docx_data.seek(0)
    return docx_data

# Streamlit UI
st.title("Lecture Transcript Analysis Agent")

# At the beginning of the Streamlit UI section, add:
if 'processed_result' not in st.session_state:
    st.session_state.processed_result = None

# In the sidebar, add the reset button:
with st.sidebar:
    st.header("Configuration")
    api_key = st.text_input("Enter your Gemini API Key:", type="password")
    if api_key:
        os.environ["GOOGLE_API_KEY"] = api_key
    
    # Add reset button
    if st.button("Reset Session"):
        st.session_state.processed_result = None
        st.rerun()
    
    st.markdown("""
    ### About
    This application processes lecture transcripts and generates a comprehensive document including:
    - Title and Speaker Information
    - 20-25 Key Quotes
    - Closing Statements
    - Briefing Document
    - Key Themes and Ideas
    - Notable Quotes
    - FAQ Section
    - Quiz Questions and Answer Key
    - Essay Questions and Answers
    - Speaker Bio
    """)

if api_key:
    # Main UI components
    speaker_name = st.text_input("Enter speaker name (lecturer):", "Alkiviadis Vazacopoulos")

    uploaded_files = st.file_uploader(
        "Upload transcript files", 
        type=['txt', 'pdf', 'docx', 'vtt'], 
        accept_multiple_files=True
    )

    if uploaded_files and speaker_name:
        if st.button("Process Transcripts"):
            processing_indicator = create_processing_indicator()
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            with processing_indicator.container():
                st.info("Processing is ongoing. This may take several minutes...")
                
                combined_text = ""
                
                # Process files
                for uploaded_file in uploaded_files:
                    with st.spinner(f"Reading {uploaded_file.name}..."):
                        try:
                            uploaded_file.seek(0)
                            file_content = read_file(uploaded_file)
                            combined_text += file_content + "\n\n"
                        except Exception as e:
                            st.error(f"Error reading {uploaded_file.name}: {str(e)}")
                            continue

                if combined_text:
                    try:
                        # Process transcript
                        processor = TranscriptProcessor(api_key)
                        st.session_state.processed_result = processor.process_transcript(
                            combined_text, 
                            speaker_name,
                            progress_bar,
                            status_text
                        )

                    except Exception as e:
                        st.error(f"An error occurred during processing: {str(e)}")
                    finally:
                        progress_bar.empty()
                        status_text.empty()

        

        # Helper function to extract date from filename
        def extract_date_from_filename(filename):
            import re
            # Look for date patterns in the filename (adjust pattern as needed)
            date_pattern = r'(\d{4}[-_]?\d{2}[-_]?\d{2})'
            match = re.search(date_pattern, filename)
            if match:
                return match.group(1).replace('_', '-')
            return time.strftime("%Y-%m-%d")  # Default to current date if no date found

        # In the display results section:
        if st.session_state.processed_result:
            st.markdown("### Processed Document")
            
            # Get the original filename and extract date
            original_filename = uploaded_files[0].name if uploaded_files else "transcript"
            file_date = extract_date_from_filename(original_filename)
            
            # Create base filename for downloads
            base_filename = f"{file_date}_Transcript_Analysis"
            
            # Create tabs for preview and raw markdown
            tab1, tab2 = st.tabs(["Preview", "Raw Markdown"])
            
            with tab1:
                cleaned_content = st.session_state.processed_result.replace("```markdown", "").replace("```", "")
                st.markdown(cleaned_content, unsafe_allow_html=True)
            
            with tab2:
                st.code(st.session_state.processed_result, language="markdown")

            # Download options
            col1, col2, col3 = st.columns(3)
            with col1:
                st.download_button(
                    label="Download as Markdown",
                    data=st.session_state.processed_result,
                    file_name=f"{base_filename}.md",
                    mime="text/markdown"
                )
            with col2:
                # Generate Word document
                docx_data = markdown_to_docx(cleaned_content)
                st.download_button(
                    label="Download as Word",
                    data=docx_data,
                    file_name=f"{base_filename}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            with col3:
                # Generate PDF
                pdf_data = markdown_to_pdf(cleaned_content, base_filename)
                st.download_button(
                    label="Download as PDF",
                    data=pdf_data,
                    file_name=f"{base_filename}.pdf",
                    mime="application/pdf"
                )
else:
    st.warning("Please enter your Gemini API key in the sidebar to continue.")
